import docx
from docx.enum.text import *
from docx.shared import Pt
import os



# Open output.txt
with open("output.txt", encoding="utf-8") as f:
    lines = [line.strip() for line in f.readlines()]
# Convert the output to a Word document, processing each line
document = docx.Document()


  
for line in lines:
    if line.startswith("# "):
        # This is a subtitle
        p = document.add_paragraph(line[2:],style='Normal')
        #p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font = p.runs[0].font
        font.name = 'Arial'
        font.size = Pt(14)
        font.bold = True
        
    elif line.startswith("$ "):
        # This is a subsubtitle
        p = document.add_paragraph(line[2:],style='Normal')
        #p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font = p.runs[0].font
        font.name = 'Arial'
        font.size = Pt(12)
        p.style = 'List Bullet' # Make it a unordered list
    else:
        # Regular content
        p = document.add_paragraph(style='Normal')
        #p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        r = p.add_run(line)
        r.font.name = 'Arial'
        r.font.size = Pt(12) 

# Save the document as output.docx
document.save('output.docx')





#abrir el docx
os.startfile("output.docx")


